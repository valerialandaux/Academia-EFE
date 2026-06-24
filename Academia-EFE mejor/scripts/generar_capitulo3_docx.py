# -*- coding: utf-8 -*-
"""Genera Capitulo3_Marco_Practico_AcademiaEFE.docx (Tecnologías Web I)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "Capitulo3_Marco_Practico_AcademiaEFE.docx"


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_h(doc, text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Arial"
    return h


def main():
    doc = Document()
    section = doc.sections[0]
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 1.5

    add_h(doc, "Capítulo 3. MARCO PRÁCTICO — DISEÑO E IMPLEMENTACIÓN DEL FRONTEND", 0)

    intro = (
        "Este capítulo documenta el análisis, la arquitectura de la información, los wireframes, "
        "el diseño visual, el diseño responsivo y la implementación técnica del frontend del sitio web "
        "de Academia EFE, un sistema orientado a la difusión de cursos de fotografía, la inscripción "
        "de interesados y la administración de solicitudes."
    )
    add_para(doc, intro)

    # 3.1
    add_h(doc, "3.1 Análisis del sitio web", 1)
    add_para(doc, "Usuarios del sistema", bold=True)
    add_para(
        doc,
        "Visitantes y potenciales alumnos: buscan información sobre la academia, el catálogo de cursos, "
        "detalle de cada taller y un canal claro para solicitar inscripción. Personal administrativo: "
        "necesita acceso restringido para revisar inscripciones, actualizar estados (pendiente, aprobada, "
        "denegada) y cerrar sesión de forma segura.",
    )
    add_para(doc, "Funciones visibles (desde la interfaz)", bold=True)
    add_para(
        doc,
        "Página de inicio con mensaje institucional, bloques de valor, cursos destacados y llamadas a la "
        "acción; sección Nosotros; catálogo de cursos con filtros y ficha de detalle; formulario de "
        "inscripción; acceso a administración con login y panel de inscripciones.",
    )
    add_para(doc, "Navegación", bold=True)
    add_para(
        doc,
        "Menú principal horizontal en escritorio (Inicio, Nosotros, Cursos, Inscribirse). En vista "
        "reducida, menú hamburguesa con panel lateral. Rutas con Vue Router; protección de /admin/"
        "inscripciones mediante token en almacenamiento local.",
    )
    add_para(doc, "Contenido", bold=True)
    add_para(
        doc,
        "Textos sobre la oferta formativa; imágenes de apoyo centralizadas en datos reutilizables; "
        "mensajes de éxito y error en formularios y en el panel administrativo.",
    )

    # 3.2
    add_h(doc, "3.2 Arquitectura de la información", 1)
    add_para(doc, "Mapa del sitio", bold=True)
    doc.add_paragraph(
        "Inicio (/) → Nosotros (/nosotros) → Cursos (/cursos) → Detalle (/cursos/:slug) → "
        "Inscripción (/inscripcion) → Admin: Login (/admin/login), Inscripciones (/admin/inscripciones)."
    )
    add_para(doc, "Estructura de navegación", bold=True)
    add_para(
        doc,
        "Menú principal sin submenús en cascada; la jerarquía Cursos–Detalle se resuelve con la ficha del "
        "curso. El título del documento (document.title) se actualiza según meta de cada ruta. "
        "Redirección a login si falta token en rutas protegidas.",
    )

    # 3.3
    add_h(doc, "3.3 Wireframes del sistema", 1)
    add_para(
        doc,
        "Figuras: insertar capturas exportadas desde Figma, Balsamiq o Draw.io. A continuación, descripción "
        "del contenido obligatorio de cada wireframe.",
    )

    add_h(doc, "3.3.1 Wireframe de la página principal", 2)
    add_para(
        doc,
        "Encabezado (logo EFE + imagen de marca), menú, hero con titular y dos botones, tres tarjetas "
        "introductorias, bloque Sobre la academia con imagen, cursos destacados, CTA final, pie de página.",
    )

    add_h(doc, "3.3.2 Wireframe del login", 2)
    add_para(
        doc,
        "Correo electrónico, contraseña, botón Entrar, enlace Volver al sitio, zona para mensaje de error.",
    )

    add_h(doc, "3.3.3 Wireframe del panel principal", 2)
    add_para(
        doc,
        "Cabecera con título, botones Actualizar y Cerrar sesión, tabla de inscripciones con columnas de "
        "datos, estado y acciones. Navegación global superior (AppNavbar).",
    )

    add_h(doc, "3.3.4 Wireframe de formularios", 2)
    add_para(
        doc,
        "Campos con etiquetas, botón de envío, mensajes de éxito y error. Inscripción: nombre completo, "
        "correo, teléfono, selector de taller.",
    )

    # 3.4
    add_h(doc, "3.4 Diseño visual del sistema", 1)

    add_h(doc, "3.4.1 Imagen corporativa", 2)
    add_para(doc, "Nombre del sistema: Academia EFE.")
    add_para(
        doc,
        "Logotipo: símbolo en /images/logo.png junto al texto EFE; colores integrados con acento cian. "
        "Eslogan alineado con la home: Captura tu visión.",
    )

    add_h(doc, "3.4.2 Tipografía", 2)
    add_para(
        doc,
        "Fuente principal: Montserrat (Google Fonts), aplicada globalmente mediante variable CSS. "
        "Pesos 300, 400, 700 y 800 para jerarquía.",
    )

    add_h(doc, "3.4.3 Paleta de colores", 2)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    hdr = ("Rol", "Color (referencia)", "Uso")
    for i, t in enumerate(hdr):
        table.rows[0].cells[i].text = t
    rows = [
        ("Primario (cian)", "#00efea / #00c4bf", "Títulos, acentos, botones, menú activo"),
        ("Fondos y paneles", "#3d3e44, #2a2b30, #32333a", "Fondo, tarjetas, paneles"),
        ("Texto secundario", "#b8b9c0", "Párrafos y descripciones"),
        ("Acción", "Degradado cian", "CTA y envío de formularios"),
        ("Alerta / error", "Tonos rojizos en mensajes", "Errores de API o validación"),
    ]
    for r, row in enumerate(rows, start=1):
        for c, cell in enumerate(row):
            table.rows[r].cells[c].text = cell

    add_h(doc, "3.4.4 Guía de estilo", 2)
    add_para(
        doc,
        "Variables CSS en :root; botones tipo píldora; secundarios con borde cian; espaciado con clamp(); "
        "radio de tarjetas 12px; sombras uniformes.",
    )

    # 3.5
    add_h(doc, "3.5 Diseño responsivo", 1)
    add_para(
        doc,
        "Computadora: menú horizontal, filtros laterales en cursos, detalle en dos columnas. Tablet: "
        "rejillas adaptativas. Celular: breakpoint 768px — menú hamburguesa, columnas apiladas, botones a "
        "ancho completo. Documentar con capturas en tres anchos.",
    )

    # 3.6
    add_h(doc, "3.6 Implementación del frontend", 1)

    add_h(doc, "3.6.1 Estructura del proyecto frontend", 2)
    doc.add_paragraph(
        "index.html, vite.config.js, package.json; src/main.js, App.vue, router/index.js, "
        "assets/style.css; components (AppNavbar, AppFooter); views (Home, Nosotros, Cursos, "
        "CursoDetalle, Inscripcion, AdminLogin, AdminInscripciones); data (cursos.js, siteImages.js); "
        "lib/api.js; public/images/."
    )

    add_h(doc, "3.6.2 Página principal", 2)
    add_para(doc, "HomePage.vue: navbar, hero, tarjetas intro, bloque nosotros, cursos destacados, CTA, footer.")

    add_h(doc, "3.6.3 Formulario de login", 2)
    add_para(doc, "AdminLoginPage.vue: correo, contraseña, botón Entrar, mensajes de error, redirección post-login.")

    add_h(doc, "3.6.4 Formulario de registro / inscripción", 2)
    add_para(
        doc,
        "El proyecto implementa inscripción pública (InscripcionPage.vue): nombre, correo, teléfono y taller. "
        "No hay registro con contraseña para alumnos; puede documentarse como mejora futura si la rúbrica "
        "lo exige.",
    )

    add_h(doc, "3.6.5 Panel principal", 2)
    add_para(doc, "AdminInscripcionesPage.vue: tabla, refresco, cierre de sesión, acciones por fila.")

    # 3.7
    add_h(doc, "3.7 Interactividad con JavaScript / Vue", 1)
    add_para(
        doc,
        "Validación HTML5 y estado de envío; mensajes dinámicos con v-if; menú interactivo con estado y "
        "watch de ruta; overlay móvil. Slider y modal genéricos no están en la versión actual.",
    )

    # 3.8
    add_h(doc, "3.8 Uso de imágenes en el sitio web", 1)
    add_para(
        doc,
        "Logotipo PNG; fotografías en hero y cursos; iconografía mínima (menú con CSS). Optimización: "
        "width/height, loading lazy/eager, object-fit cover.",
    )

    # 3.9
    add_h(doc, "3.9 Accesibilidad web", 1)
    add_para(
        doc,
        "Texto alternativo en imágenes relevantes; contraste sobre fondos oscuros; foco visible en inputs; "
        "label asociado a id; aria-label, aria-expanded, role alert/status; lang=es; prefers-reduced-motion.",
    )

    # 3.10
    add_h(doc, "3.10 Pruebas del frontend", 1)
    add_para(
        doc,
        "Visualización, responsividad, usabilidad (flujos inscripción y admin), compatibilidad en Chrome, "
        "Firefox y Edge. Registrar resultados en tabla anexa si aplica.",
    )

    # 3.11
    add_h(doc, "3.11 Publicación del sitio web", 1)
    add_para(
        doc,
        "Servidor local: npm run dev; API con npm run server o npm run dev:full. Hosting: Netlify o "
        "GitHub Pages sobre build estático (dist). Repositorio Git con README.",
    )

    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 1.5
        for r in p.runs:
            if r.font.name is None or r.font.name == "":
                r.font.name = "Arial"
            if r.font.size is None:
                r.font.size = Pt(12)

    doc.save(OUT)
    print(f"Guardado: {OUT}")


if __name__ == "__main__":
    main()
